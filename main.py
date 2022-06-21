import tkinter as tk
from docx import Document
from DataFrm.DataFrm import DateFrm


def get_info():

    """
    1. 抓取Entry資料並寫入DateFrm的dict內
    2. 產生.docx 文件檔 並寫入使用者所填入的資料 並在開頭加上一張公司LOGO圖片
    """

    DateFrm['chineseName_var'] = chineseNameEntry.get()
    DateFrm['englishName_var'] = englishNameEntry.get()
    DateFrm['emailAddress_var'] = emailAddressEntry.get()
    DateFrm['companyPhone_var'] = companyPhoneEntry.get()
    DateFrm['mobilePhone_var'] = mobilePhoneEntry.get()

    docx = Document()

    docx.add_picture('Image/INNOLIGHT.PNG')

    docx.add_paragraph(DateFrm['chineseName_var'], style='List Bullet')
    docx.add_paragraph(DateFrm['englishName_var'], style='List Bullet')
    docx.add_paragraph(DateFrm['emailAddress_var'], style='List Bullet')
    docx.add_paragraph(DateFrm['companyPhone_var'], style='List Bullet')
    docx.add_paragraph(DateFrm['mobilePhone_var'], style='List Bullet')

    docx.save('OutlookSign.docx')

    print(DateFrm['chineseName_var'],
          DateFrm['englishName_var'],
          DateFrm['emailAddress_var'],
          DateFrm['companyPhone_var'],
          DateFrm['mobilePhone_var']
          )


# 產生主視窗 設定大小 以及抬頭
root = tk.Tk()
root.geometry('350x150')
root.title('Create outlook sign')

# 輸入視窗資料型別定義
chineseName_var = tk.StringVar()
englishName_var = tk.StringVar()
emailAddress_var = tk.StringVar()
companyPhone_var = tk.StringVar()
mobilePhone_var = tk.StringVar()

# 創建標示標籤
chineseNameLabel = tk.Label(root, text='chineseName', font=('caliber', 10, 'bold'))
englishNameLabel = tk.Label(root, text='englishName', font=('caliber', 10, 'bold'))
emailAddressLabel = tk.Label(root, text='emailAddress', font=('caliber', 10, 'bold'))
companyPhoneLabel = tk.Label(root, text='companyPhone', font=('caliber', 10, 'bold'))
mobilePhoneLabel = tk.Label(root, text='mobilePhone', font=('caliber', 10, 'bold'))

# 創建輸入視窗
chineseNameEntry = tk.Entry(root, textvariable=chineseName_var, font=('caliber', 10, 'bold'))
englishNameEntry = tk.Entry(root, textvariable=englishName_var, font=('caliber', 10, 'bold'))
emailAddressEntry = tk.Entry(root, textvariable=emailAddress_var, font=('caliber', 10, 'bold'))
companyPhoneEntry = tk.Entry(root, textvariable=companyPhone_var, font=('caliber', 10, 'bold'))
mobilePhoneEntry = tk.Entry(root, textvariable=mobilePhone_var, font=('caliber', 10, 'bold'))
createButton = tk.Button(root, text='CreateSign', command=get_info, font=('caliber', 10, 'bold'))

# 定位點
chineseNameLabel.grid(row=0, column=0)
chineseNameEntry.grid(row=0, column=1)
englishNameLabel.grid(row=1, column=0)
englishNameEntry.grid(row=1, column=1)
emailAddressLabel.grid(row=2, column=0)
emailAddressEntry.grid(row=2, column=1)
companyPhoneLabel.grid(row=3, column=0)
companyPhoneEntry.grid(row=3, column=1)
mobilePhoneLabel.grid(row=4, column=0)
mobilePhoneEntry.grid(row=4, column=1)
createButton.grid(row=5, column=3)

root.mainloop()
